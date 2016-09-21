# -*- coding: utf-8 -*-
from docx import Document
import csv
import pudb
import string

doc = Document('SICB2016Posters.docx')

# bold_texts = [[run.text for run in p.runs if run.bold] for p in doc.paragraphs]


printable = set(string.printable)
# p = doc.paragraphs[0]

email_title_institution_fname_lname_abstract = []
ps = doc.paragraphs

for i, p in enumerate(ps):
    if '@' in p.text:

        # get email
        text = p.text
        text_list = text.split(" ")
        for word in text_list:
            if '@' in word:
                email =  filter(lambda x: x in printable, word)

        # get title
        title = ps[i+1].text

        # get institution
        if ";" in p.text:
            if p.text.count(";") > 1:
                split = p.text.split(";")


                for idx, el in enumerate(split):
                    if email in el:
                        in_all = split[idx-1]

                # in_all = split[-2]
            else:
                in_2 = p.text.split(";")[0]
                in_1 = ps[i-1].text.split(";")[-1]
                in_all = " ".join((in_1, in_2))
        else:
            try:
                in_all = ps[i-1].text.split(";")[-2]
            except:
                in_all = ""

        if len(in_all) > 300:

            

            k = 0
            j = 0
            first = False

            l = list(text)
            at_index = l.index("@")
            for i in reversed(range(at_index)):
                if l[i] == ";":
                    if first is False:
                        
                        k = i
                        first = True
                    else:
                        
                        j = i
            
            in_all = text[j:k]

        if email == "Roger.Anderson@wwu.edu":
            pu.db

        # get Name 
        m = i
        prev = ps[m]
        flag = False
        while "*" not in prev.text:
            m -= 1
            prev = ps[m]
            if prev.text == "" or prev.text == "\n":
                prev = ps[m+1]
                flag = True
                break

        if flag is False:
            name = prev.text.split("*")
            name = name[0].split(",")
        else:
            name = prev.text.split(";")[0]
            name = name.split(",")

        try:
            fname = name[1]
        except:
            fname = name[0]
        try:
            lname = name[0].split(" ")[1]
        except:
            lname = name[0]
            if lname == fname:
                lname = ""

        # get abstract
        abstract = ps[i+2].text
        if len(abstract) < 200:
            abstract = ps[i+1].text
            if len(abstract) < 200:
                abstract = p.text
                if len(abstract) < 200:
                    abstract = ps[i+3].text
                    if len(abstract) < 200:

                        n = i + 2
                        abstract = ""
                        while len(abstract) < 200:
                            abstract += ps[n].text
                            n += 1
                        if len(abstract) < 200:
                            pu.db


        out = [email, title, in_all, fname, lname, abstract]
        out = [filter(lambda x: x in printable, word) for word in out]

        # fnames = [el[3] for el in out]
        # lnames = [el[4] for el in out]

        # for el in out:
        #     fname = el[3]
        #     lname = el[4]
        #     count = 0
        #     for zel in out:
        #         if fname == zel[3] and lname == zel[4]:
        #             matches = [[x, el] for x, el in enumerate(out) if el[3] == fname and el[4] == lname]
        #             for y, match in enumerate(matches):
        #                 if y != 0:
        #                     del out[match[0]]


        

        email_title_institution_fname_lname_abstract.append(out)

with open('sicb_p.csv', 'w') as mycsvfile:
    thedatawriter = csv.writer(mycsvfile)
    for el in email_title_institution_fname_lname_abstract:
        thedatawriter.writerow(el)